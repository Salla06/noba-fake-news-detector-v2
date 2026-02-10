"""
Utility functions for FCC Fake News Detector - IMPROVED VERSION
"""

import re
import requests
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from langdetect import detect

def get_language_name(code):
    """Obtenir le nom de la langue"""
    languages = {
        'auto': '🌐 Auto-detect',
        'en': '🇬🇧 English',
        'fr': '🇫🇷 Français',
        'es': '🇪🇸 Español',
        'ar': '🇸🇦 العربية',
        'zh-CN': '🇨🇳 中文'
    }
    return languages.get(code, code)

def detect_language(text):
    """Détecter la langue"""
    try:
        return detect(text)
    except:
        return 'en'

def translate_to_english(text, source_lang=None):
    """Traduire vers l'anglais et retourner le texte original aussi"""
    try:
        if source_lang is None or source_lang == 'auto':
            detected = detect(text)
        else:
            detected = source_lang
        
        if detected == 'en':
            return text, 'en', None  # Pas de traduction nécessaire
        
        translator = GoogleTranslator(source=detected, target='en')
        translated = translator.translate(text)
        
        return translated, detected, text  # Retourner le texte original
    except Exception as e:
        return text, 'en', None

def extract_text_from_url(url):
    """Extraire le texte d'une URL (page web ou document)"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        # PDF
        if 'application/pdf' in response.headers.get('Content-Type', ''):
            import PyPDF2
            import io
            pdf_file = io.BytesIO(response.content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        
        # Word
        elif 'application/vnd.openxmlformats' in response.headers.get('Content-Type', ''):
            import docx
            import io
            doc_file = io.BytesIO(response.content)
            doc = docx.Document(doc_file)
            return "\n".join([p.text for p in doc.paragraphs])
        
        # HTML
        else:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
                element.decompose()
            
            text = soup.get_text(separator='\n')
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            return text
    
    except Exception as e:
        raise Exception(f"Erreur extraction: {str(e)}")

def extract_text_from_file(uploaded_file):
    """Extraire texte d'un fichier uploadé"""
    try:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if file_ext == 'txt':
            return uploaded_file.read().decode('utf-8')
        
        elif file_ext == 'pdf':
            import PyPDF2
            pdf = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf.pages:
                text += page.extract_text()
            return text
        
        elif file_ext in ['docx', 'doc']:
            import docx
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        
        elif file_ext in ['xlsx', 'xls']:
            import pandas as pd
            df = pd.read_excel(uploaded_file)
            return df.to_string()
        
        return None
    except:
        return None
